import os
from pydantic import Extra
import requests
from typing import Any, List, Mapping, Optional
import json
import ast
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain.llms.base import LLM
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain.prompts import PromptTemplate
# Run chain
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS # type: ignore
from langchain_ollama import OllamaEmbeddings # type: ignore
import warnings
warnings.filterwarnings('ignore')
 
##################################
 
os.environ["token"] = "Bearer token|21031658-00bd-47ae-bbcb-c4d9e6c1146f|c87e47bd3697e43906884735e46c827bfc7dee361628b9e7ad26d369bfadc098"
token = os.environ["token"]
 
#################################
 
from typing import ClassVar
 
parser = StrOutputParser()
 
class LlamaLLM(LLM):
    llm_url: ClassVar[str] = 'https://api.lab45.ai/v1.1/skills/completion/query'
   
    backend:        Optional[str]   = 'gpt-35-turbo-16k'
    temp:           Optional[float] = 0.7
    top_p:          Optional[float] = 0.1
    top_k:          Optional[int]   = 40
    n_batch:        Optional[int]   = 8
    n_threads:      Optional[int]   = 4
    n_predict:      Optional[int]   = 256
    max_tokens:     Optional[int]   = 256
    repeat_last_n:  Optional[int]   = 64
    repeat_penalty: Optional[float] = 1.18
 
    class Config:
        extra = Extra.forbid
 
    @property
    def _llm_type(self) -> str:
        return "gpt-35-turbo-16k"
   
    @property
    def _get_model_default_parameters(self):
        return {
            "max_tokens": self.max_tokens,
            #"n_predict": self.n_predict,
            "top_k": self.top_k,
            "top_p": self.top_p,
            "temperature": self.temp,
            #"n_batch": self.n_batch,
            #"repeat_penalty": self.repeat_penalty,
            #"repeat_last_n": self.repeat_last_n,
        }
 
    def _call(
        self,
        prompt: str,
        user: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        if stop is not None:
            raise ValueError("stop kwargs are not permitted.")
 
        payload = {
        "messages": [
            {
            "content": prompt,
            "role": user
            }
        ],
        "skill_parameters": {
            "model_name": "gpt-35-turbo-16k",
            "max_output_tokens": 4096,
            "temperature": 0,
            "top_k": 5
        },
        "stream_response": False
        }
 
        headers = {"Content-Type": "application/json","Authorization": token}
 
        response = requests.post(self.llm_url, json=payload, headers=headers, verify=False)
 
       # print("API Response:", response.json())
        response.raise_for_status()
 
        return response.json()  # get the response from the API
 
    @property
    def _identifying_params(self) -> Mapping[str, Any]:
        """Get the identifying parameters."""
        return {
            "llmUrl": self.llm_url,
            'model_parameters': self._get_model_default_parameters
           
            }
   
################################################################################################
 
 
from pydantic import BaseModel
from typing import Dict
 
class InterviewQA(BaseModel):
    questions_and_answers: Dict[str, str]
 
####################################################PYDANTIC OUTPUT PARSER
   
llm = LlamaLLM()
user="user"
 
import streamlit as st
import fitz  # PyMuPDF
import docx
import re
 
# ---------------------- File Reading Functions ----------------------
 
def read_pdf(file):
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in pdf_document:
        text += page.get_text()
    return text
 
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text
 
def extract_text_from_txt(file):
    return file.read().decode('utf-8')
 
# ---------------------- LLM Interaction Functions ----------------------
 
def analyze_jd_resume_match(job_description, resume):
    prompt = f"""
Compare the following Job Description and Candidate Resume. Provide a detailed summary including:
 
1. Technical Skill Match
2. Functional Skill Match
3. Experience Relevance
4. Domain Knowledge Alignment
5. Soft Skills Match
6. Overall Match Percentage between 1 to 100, give exact percentage
 
Job Description:
{job_description}
 
Candidate Resume:
{resume}
"""
    response = llm._call(prompt, user)
    content = response['data']['content']
    match = re.search(r"(\d+)%", content)
    score = int(match.group(1)) if match else 0
    return content, score
 
def generate_interview_questions(job_description):
    prompt = f"""
Given the following Job Description: {job_description}
 
Generate a structured set of 100 highly technical interview questions **with detailed answers**. Focus on:
 
- Core technical concepts
- System design and architecture
- Algorithms and data structures
- Language-specific coding challenges
- Debugging and optimization
- Real-world problem-solving scenarios
- Domain-specific tools, frameworks, and methodologies
 
Avoid generic or behavioral questions. Return the output as a JSON object that conforms to the following Pydantic model:
 
class InterviewQA(BaseModel):
    questions_and_answers: Dict[str, str]
"""
    response = llm._call(prompt, user)
    return response['data']['content']
 
def get_next_unasked_question():
    for q, a in st.session_state.questions:
        if q not in st.session_state.asked_questions:
            st.session_state.asked_questions.add(q)
            return q, a
    return None, None
 
def simulate_interview(job_description, resume, responses):
    prompt = f"Job Description:\n{job_description}\n\nCandidate Resume:\n{resume}\n\nInterview Responses:\n{responses}\n\nAnalyze the candidate's skills based on the interview responses."
    response = llm._call(prompt, user)
    return response['data']['content']
 
def evaluate_interview_transcript(transcript: str):
    prompt = f"""
Evaluate the following interview transcript for a candidate. Assess the candidate across four key dimensions relevant to the role. For each category, provide a score out of 10, a brief justification, and then summarize the candidate's overall strengths, areas for improvement, and a final recommendation (e.g., Strong Hire, Hire, Weak Hire, No Hire).
 
Evaluation Categories:
 
1. Technical Capability
   - Problem-solving and coding skills
   - Technical execution and practical application
   - Design/technical depth and rationale
 
2. Domain Expertise
   - Depth of knowledge in the relevant domain (e.g., UX, IoT, Data Science, etc.)
   - Familiarity with tools, frameworks, and methodologies
 
3. Communication & Collaboration
   - Clarity of expression and articulation of thoughts
   - Responsiveness to interviewer prompts and engagement
   - Ability to work with cross-functional teams and stakeholder awareness
 
4. Depth & Impact
   - Level of detail in describing past projects and decisions
   - Demonstrated outcomes, metrics, and business/user impact
 
Transcript:
{transcript}
"""
    response = llm._call(prompt, user)
    return response['data']['content']
 
# ---------------------- Streamlit App ----------------------
 
st.title("Interactive Interview Simulator")
 
# Initialize session state variables
default_session_state = {
    "interview_stage": 0,
    "questions": [],
    "responses": [],
    "allow_interview": False,
    "start_interview": False,
    "match_score": None,
    "match_summary": None,
    "question_index": 0,
    "chat_history": [],
    "chat_mode_active": False,
    "asked_questions": set(),
    "depth_level": 1,
}
 
 
for key, value in default_session_state.items():
    if key not in st.session_state or st.session_state[key] is None:
        st.session_state[key] = value
 
 
# Upload job description
job_description_file = st.file_uploader("Upload Job Description (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
if job_description_file:
    if job_description_file.type == "application/pdf":
        job_description = read_pdf(job_description_file)
    elif job_description_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        job_description = extract_text_from_docx(job_description_file)
    elif job_description_file.type == "text/plain":
        job_description = extract_text_from_txt(job_description_file)
    st.text_area("Job Description", job_description, height=200)
else:
    job_description = None
 
# Upload candidate resume
resume_file = st.file_uploader("Upload Candidate Resume (PDF, DOCX)", type=["pdf", "docx"])
if resume_file:
    if resume_file.type == "application/pdf":
        resume = read_pdf(resume_file)
    elif resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume = extract_text_from_docx(resume_file)
    elif resume_file.type == "text/plain":
        resume = extract_text_from_txt(resume_file)
    st.text_area("Candidate Resume", resume, height=200)
else:
    resume = None
 
# JD-Resume Matching Analysis
if job_description and resume:
    if st.button("Analyze JD-Resume Match"):
        match_summary, match_score = analyze_jd_resume_match(job_description, resume)
        st.session_state.match_summary = match_summary
        st.session_state.match_score = match_score
        st.session_state.start_interview = False  # Reset interview flag
 
# Display match results
if st.session_state.match_score is not None:
    st.subheader("JD-Resume Matching Summary")
    st.text_area("Matching Analysis", st.session_state.match_summary, height=300)
 
    if st.session_state.match_score < 60:
        st.warning(f"Overall Match Score is {st.session_state.match_score}%. This is considered low.")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Continue and Apply Anyway"):
                st.session_state.allow_interview = True
                st.session_state.start_interview = True
        with col2:
            if st.button("Skip This Role"):
                st.session_state.allow_interview = False
                st.session_state.start_interview = False
    else:
        st.success(f"Overall Match Score is {st.session_state.match_score}%. You are a good fit!")
        st.session_state.allow_interview = True
        st.session_state.start_interview = True
 
def safe_parse_json(raw_json: str):
    # Remove Markdown formatting
    cleaned = re.sub(r"^```json|```$", "", raw_json.strip(), flags=re.MULTILINE).strip()
 
    # Try standard JSON parsing
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
 
    # Try literal_eval
    try:
        return ast.literal_eval(cleaned)
    except (ValueError, SyntaxError):
        pass
 
    # Try fixing common issues
    cleaned = cleaned.replace('\n', '\\n').replace('\r', '')
    cleaned = re.sub(r'(?<!\\)"', '\\"', cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None
 
# Generate Interview Questions
if st.session_state.start_interview and st.session_state.allow_interview and not st.session_state.questions:
 
    raw_json = generate_interview_questions(job_description)
    parsed = safe_parse_json(raw_json)
 
    if parsed and isinstance(parsed, dict):
        qa_dict = {
            q.strip(): a.strip()
            for q, a in parsed.items()
            if isinstance(q, str) and isinstance(a, str)
        }
        st.session_state.questions = [(q, a) for q, a in qa_dict.items()]
        st.session_state.interview_stage = 0
        st.session_state.responses = []
        st.session_state.start_interview = False
    else:
        st.error("Failed to parse interview questions. Please try again.")
 
# ---------------------- Chat Mode Interview ----------------------
 
st.title("Real-Time Chat Interview Simulator")
 
# Initialize chat session state
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
 
if "chat_mode_active" not in st.session_state:
    st.session_state.chat_mode_active = False
 
raw_json=[]
parsed=[]
 
# Start Chat Interview
if job_description and resume and st.button("Start Chat Interview"):
    st.session_state.chat_mode_active = True
    st.session_state.chat_history = [{
        "role": "interviewer",
        "content": "Hello! Let's begin your interview based on the job description. Can you briefly introduce yourself?"
    }]
 
    # Use existing questions if available
    if st.session_state.questions:
        parsed = {q: a for q, a in st.session_state.questions}
    else:
        with st.spinner("Generating interview questions..."):
            raw_json = generate_interview_questions(job_description)
            parsed = safe_parse_json(raw_json)
            st.write(raw_json)
            if parsed and isinstance(parsed, dict):
                qa_dict = {
                    q.strip(): a.strip()
                    for q, a in parsed.items()
                    if isinstance(q, str) and isinstance(a, str)
                }
               
                # Limit to 5 questions
                limited_questions = list(qa_dict.items())[:5]
                st.session_state.questions = limited_questions
               
            else:
                st.error("Failed to parse interview questions. Please try again.")
                st.session_state.chat_mode_active = False
 
if st.session_state.chat_mode_active:
    current_question, current_answer = get_next_unasked_question()
 
    # Display chat history
    st.subheader("Interview Chat")
 
    for msg in st.session_state.chat_history:
        if msg["role"] == "interviewer":
            st.markdown(f"**🧑‍💼 Interviewer:** {msg['content']}")
        else:
            st.markdown(f"**🙋 You:** {msg['content']}")
 
    # Input box for candidate response
    if "temp_input" not in st.session_state:
        st.session_state.temp_input = ""
 
    user_input = st.text_input("Your response:", value=st.session_state.temp_input, key="chat_input")
 
    # Send button logic
    if st.button("Submit") and user_input and not st.session_state.get("response_sent", False):
        st.session_state.response_sent = True
        # Add candidate response to chat history
        st.session_state.chat_history.append({"role": "candidate", "content": user_input})
 
        last_response = next(
            (msg['content'] for msg in reversed(st.session_state.chat_history) if msg['role'] == 'candidate'),
            ""
        )
 
        chat_prompt = "\n".join(
            f"{'Candidate' if msg['role']=='candidate' else 'Interviewer'}: {msg['content']}"
            for msg in st.session_state.chat_history
        )
 
        full_prompt = f"""
You are a senior technical interviewer conducting a live interview.
 
Job Description:
{job_description}
 
Candidate Resume:
{resume}
 
Interview Depth Level: {st.session_state.depth_level}
 
Instructions:
- Conduct the interview in a technical and progressive manner.
- At depth level 1, ask basic questions.
- At depth level 2-3, ask intermediate questions.
- At depth level 4+, ask advanced or domain-specific questions.
- Avoid repeating previous questions.
- Use the candidate's resume to tailor questions.
 
Conversation so far:
{chat_prompt}
 
Last Candidate Response:
{last_response}
 
Continue the interview with a new technical question or follow-up.
"""
 
        # Get LLM response
        response = llm._call(full_prompt, user)
        reply = response['data']['content']
 
        # Add interviewer response to history
        st.session_state.chat_history.append({"role": "interviewer", "content": reply})
 
        # Advance depth level
        st.session_state.depth_level += 1
 
        # Advance question index
        st.session_state.question_index += 1
 
        # Clear input safely
        st.session_state.temp_input = ""
        st.session_state.response_sent = False
        st.rerun()
 
    # Evaluation option after all questions
    if st.session_state.question_index >= len(st.session_state.questions):
        full_chat = "\n".join(
            f"{msg['role'].capitalize()}: {msg['content']}"
            for msg in st.session_state.chat_history
        )
 
        if st.button("Evaluate Interview"):
            with st.spinner("Evaluating interview..."):
                evaluation = evaluate_interview_transcript(full_chat)
                st.subheader("📊 Interview Evaluation")
                st.text_area("Evaluation Summary", evaluation, height=400)
 
    # Reset chat option
    if st.button("Reset Chat"):
        st.session_state.chat_history = []
        st.session_state.chat_mode_active = False
        st.session_state.temp_input = ""
        st.session_state.response_sent = False
        st.session_state.depth_level = 1
        st.session_state.question_index = 0
        st.session_state.asked_questions = set()
        st.rerun()
 
 
 
import os
from pydantic import Extra
import requests
from typing import Any, List, Mapping, Optional
import json
import ast
from langchain.callbacks.manager import CallbackManagerForLLMRun
from langchain.llms.base import LLM
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain.prompts import PromptTemplate
# Run chain
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS # type: ignore
from langchain_ollama import OllamaEmbeddings # type: ignore
import warnings
warnings.filterwarnings('ignore')
 
##################################
 
os.environ["token"] = "Bearer token|21031658-00bd-47ae-bbcb-c4d9e6c1146f|c87e47bd3697e43906884735e46c827bfc7dee361628b9e7ad26d369bfadc098"
token = os.environ["token"]
 
#################################
 
from typing import ClassVar
 
parser = StrOutputParser()
 
class LlamaLLM(LLM):
    llm_url: ClassVar[str] = 'https://api.lab45.ai/v1.1/skills/completion/query'
   
    backend:        Optional[str]   = 'gpt-35-turbo-16k'
    temp:           Optional[float] = 0.7
    top_p:          Optional[float] = 0.1
    top_k:          Optional[int]   = 40
    n_batch:        Optional[int]   = 8
    n_threads:      Optional[int]   = 4
    n_predict:      Optional[int]   = 256
    max_tokens:     Optional[int]   = 256
    repeat_last_n:  Optional[int]   = 64
    repeat_penalty: Optional[float] = 1.18
 
    class Config:
        extra = Extra.forbid
 
    @property
    def _llm_type(self) -> str:
        return "gpt-35-turbo-16k"
   
    @property
    def _get_model_default_parameters(self):
        return {
            "max_tokens": self.max_tokens,
            #"n_predict": self.n_predict,
            "top_k": self.top_k,
            "top_p": self.top_p,
            "temperature": self.temp,
            #"n_batch": self.n_batch,
            #"repeat_penalty": self.repeat_penalty,
            #"repeat_last_n": self.repeat_last_n,
        }
 
    def _call(
        self,
        prompt: str,
        user: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        if stop is not None:
            raise ValueError("stop kwargs are not permitted.")
 
        payload = {
        "messages": [
            {
            "content": prompt,
            "role": user
            }
        ],
        "skill_parameters": {
            "model_name": "gpt-35-turbo-16k",
            "max_output_tokens": 4096,
            "temperature": 0,
            "top_k": 5
        },
        "stream_response": False
        }
 
        headers = {"Content-Type": "application/json","Authorization": token}
 
        response = requests.post(self.llm_url, json=payload, headers=headers, verify=False)
 
       # print("API Response:", response.json())
        response.raise_for_status()
 
        return response.json()  # get the response from the API
 
    @property
    def _identifying_params(self) -> Mapping[str, Any]:
        """Get the identifying parameters."""
        return {
            "llmUrl": self.llm_url,
            'model_parameters': self._get_model_default_parameters
           
            }
   
################################################################################################
 
 
from pydantic import BaseModel
from typing import Dict
 
class InterviewQA(BaseModel):
    questions_and_answers: Dict[str, str]
 
####################################################PYDANTIC OUTPUT PARSER
   
llm = LlamaLLM()
user="user"
 
import streamlit as st
import fitz  # PyMuPDF
import docx
import re
 
# ---------------------- File Reading Functions ----------------------
 
def read_pdf(file):
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in pdf_document:
        text += page.get_text() # type: ignore
    return text
 
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text
 
def extract_text_from_txt(file):
    return file.read().decode('utf-8')
 
# ---------------------- LLM Interaction Functions ----------------------
 
def analyze_jd_resume_match(job_description, resume):
    prompt = f"""
Compare the following Job Description and Candidate Resume. Provide a detailed summary including:
 
1. Technical Skill Match
2. Functional Skill Match
3. Experience Relevance
4. Domain Knowledge Alignment
5. Soft Skills Match
6. Overall Match Percentage between 1 to 100, give exact percentage
 
Job Description:
{job_description}
 
Candidate Resume:
{resume}
"""
    response = llm._call(prompt, user)
    content = response['data']['content']
    match = re.search(r"(\d+)%", content)
    score = int(match.group(1)) if match else 0
    return content, score
 
def generate_interview_questions(job_description):
    prompt = f"""
Given the following Job Description: {job_description}
 
Generate a structured set of 100 highly technical interview questions **with detailed answers**. Focus on:
 
- Core technical concepts
- System design and architecture
- Algorithms and data structures
- Language-specific coding challenges
- Debugging and optimization
- Real-world problem-solving scenarios
- Domain-specific tools, frameworks, and methodologies
 
Avoid generic or behavioral questions. Return the output as a JSON object that conforms to the following Pydantic model:
 
class InterviewQA(BaseModel):
    questions_and_answers: Dict[str, str]
"""
    response = llm._call(prompt, user)
    return response['data']['content']
 
 
def get_next_unasked_question():
    for q, a in st.session_state.questions:
        if q not in st.session_state.asked_questions:
            st.session_state.asked_questions.add(q)
            return q, a
    return None, None
 
def simulate_interview(job_description, resume, responses):
    prompt = f"Job Description:\n{job_description}\n\nCandidate Resume:\n{resume}\n\nInterview Responses:\n{responses}\n\nAnalyze the candidate's skills based on the interview responses."
    response = llm._call(prompt, user)
    return response['data']['content']
 
def evaluate_interview_transcript(transcript: str):
    prompt = f"""
Evaluate the following interview transcript for a candidate. Assess the candidate across four key dimensions relevant to the role. For each category, provide a score out of 10, a brief justification, and then summarize the candidate's overall strengths, areas for improvement, and a final recommendation (e.g., Strong Hire, Hire, Weak Hire, No Hire).
 
Evaluation Categories:
 
1. Technical Capability
   - Problem-solving and coding skills
   - Technical execution and practical application
   - Design/technical depth and rationale
 
2. Domain Expertise
   - Depth of knowledge in the relevant domain (e.g., UX, IoT, Data Science, etc.)
   - Familiarity with tools, frameworks, and methodologies
 
3. Communication & Collaboration
   - Clarity of expression and articulation of thoughts
   - Responsiveness to interviewer prompts and engagement
   - Ability to work with cross-functional teams and stakeholder awareness
 
4. Depth & Impact
   - Level of detail in describing past projects and decisions
   - Demonstrated outcomes, metrics, and business/user impact
 
Transcript:
{transcript}
"""
    response = llm._call(prompt, user)
    return response['data']['content'] # type: ignore
 
# ---------------------- Streamlit App ----------------------
 
st.title("Interactive Interview Simulator")
 
# Initialize session state variables
default_session_state = {
    "interview_stage": 0,
    "questions": [],
    "responses": [],
    "allow_interview": False,
    "start_interview": False,
    "match_score": None,
    "match_summary": None,
    "question_index": 0,
    "chat_history": [],
    "chat_mode_active": False,
    "asked_questions": set(),
    "depth_level": 1,
}
 
for key, value in default_session_state.items():
    if key not in st.session_state or st.session_state[key] is None:
        st.session_state[key] = value
 
 
# Upload job description
job_description_file = st.file_uploader("Upload Job Description (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
if job_description_file:
    if job_description_file.type == "application/pdf":
        job_description = read_pdf(job_description_file)
    elif job_description_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        job_description = extract_text_from_docx(job_description_file)
    elif job_description_file.type == "text/plain":
        job_description = extract_text_from_txt(job_description_file)
    st.text_area("Job Description", job_description, height=200)
else:
    job_description = None
 
# Upload candidate resume
resume_file = st.file_uploader("Upload Candidate Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
if resume_file:
    if resume_file.type == "application/pdf":
        resume = read_pdf(resume_file)
    elif resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume = extract_text_from_docx(resume_file)
    elif resume_file.type == "text/plain":
        resume = extract_text_from_txt(resume_file)
    st.text_area("Candidate Resume", resume, height=200)
else:
    resume = None
 
# JD-Resume Matching Analysis
if job_description and resume:
    if st.button("Analyze JD-Resume Match"):
        match_summary, match_score = analyze_jd_resume_match(job_description, resume)
        st.session_state.match_summary = match_summary
        st.session_state.match_score = match_score
        st.session_state.start_interview = False  # Reset interview flag
 
# Display match results
if st.session_state.match_score is not None:
    st.subheader("JD-Resume Matching Summary")
    st.text_area("Matching Analysis", st.session_state.match_summary, height=300)
 
    if st.session_state.match_score < 60:
        st.warning(f"Overall Match Score is {st.session_state.match_score}%. This is considered low.")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Continue and Apply Anyway"):
                st.session_state.allow_interview = True
                st.session_state.start_interview = True
        with col2:
            if st.button("Skip This Role"):
                st.session_state.allow_interview = False
                st.session_state.start_interview = False
    else:
        st.success(f"Overall Match Score is {st.session_state.match_score}%. You are a good fit!")
        st.session_state.allow_interview = True
        st.session_state.start_interview = True
 
def safe_parse_json(raw_json: str):
    # Remove Markdown formatting
    cleaned = re.sub(r"^```json|```$", "", raw_json.strip(), flags=re.MULTILINE).strip()
 
    # Try standard JSON parsing
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
 
    # Try literal_eval
    try:
        return ast.literal_eval(cleaned)
    except (ValueError, SyntaxError):
        pass
 
    # Try fixing common issues
    cleaned = cleaned.replace('\n', '\\n').replace('\r', '')
    cleaned = re.sub(r'(?<!\\)"', '\\"', cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None
 
# Generate Interview Questions
if st.session_state.start_interview and st.session_state.allow_interview and not st.session_state.questions:
 
    raw_json = generate_interview_questions(job_description)
    parsed = safe_parse_json(raw_json)
 
    if parsed and isinstance(parsed, dict):
        qa_dict = {
            q.strip(): a.strip()
            for q, a in parsed.items()
            if isinstance(q, str) and isinstance(a, str)
        }
        st.session_state.questions = [(q, a) for q, a in qa_dict.items()]
        st.session_state.interview_stage = 0
        st.session_state.responses = []
        st.session_state.start_interview = False
    else:
        st.error("Failed to parse interview questions. Please try again.")
 
# ---------------------- Chat Mode Interview ----------------------
 
st.title("Real-Time Chat Interview Simulator")
 
# Initialize chat session state
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
 
if "chat_mode_active" not in st.session_state:
    st.session_state.chat_mode_active = False
 
raw_json=[]
parsed=[]
question=5
 
# Start Chat Interview
if job_description and resume and st.button("Start Chat Interview"):
    st.session_state.chat_mode_active = True
    st.session_state.chat_history = [{
        "role": "interviewer",
        "content": "Hello! Let's begin your interview based on the job description. Can you briefly introduce yourself?"
    }]
 
    # Use existing questions if available
    if st.session_state.questions:
        parsed = {q: a for q, a in st.session_state.questions}
    else:
        with st.spinner("Generating interview questions..."):
            raw_json = generate_interview_questions(job_description)
            parsed = safe_parse_json(raw_json)
            st.write(raw_json)
            if parsed and isinstance(parsed, dict):
                qa_dict = {
                    q.strip(): a.strip()
                    for q, a in parsed.items()
                    if isinstance(q, str) and isinstance(a, str)
                }
               
                # Limit to 5 questions
                limited_questions = list(qa_dict.items())[:question]
                st.session_state.questions = limited_questions
               
            else:
                st.error("Failed to parse interview questions. Please try again.")
                st.session_state.chat_mode_active = False
 
if st.session_state.chat_mode_active:
    current_question, current_answer = get_next_unasked_question()
 
    # Display chat history
    st.subheader("Interview Chat")
 
    for msg in st.session_state.chat_history:
        if msg["role"] == "interviewer":
            st.markdown(f"**🧑‍💼 Interviewer:** {msg['content']}")
        else:
            st.markdown(f"**🙋 You:** {msg['content']}")
 
    # Input box for candidate response
    if "temp_input" not in st.session_state:
        st.session_state.temp_input = ""
 
    user_input = st.text_input("Your response:", value=st.session_state.temp_input, key="chat_input")
 
    # Send button logic
    if st.button("Submit") and user_input and not st.session_state.get("response_sent", False):
        st.session_state.response_sent = True
        # Add candidate response to chat history
        st.session_state.chat_history.append({"role": "candidate", "content": user_input})
 
        last_response = next(
            (msg['content'] for msg in reversed(st.session_state.chat_history) if msg['role'] == 'candidate'),
            ""
        )
 
        chat_prompt = "\n".join(
            f"{'Candidate' if msg['role']=='candidate' else 'Interviewer'}: {msg['content']}"
            for msg in st.session_state.chat_history
        )
 
        full_prompt = f"""
You are a senior technical interviewer conducting a live interview.
 
Job Description:
{job_description}
 
Candidate Resume:
{resume}
 
Interview Depth Level: {st.session_state.depth_level}
 
Instructions:
- Conduct the interview in a technical and progressive manner.
- At depth level 1, ask basic questions.
- At depth level 2-3, ask intermediate questions.
- At depth level 4+, ask advanced or domain-specific questions.
- Avoid repeating previous questions.
- Use the candidate's resume to tailor questions.
- Ensure that the candidate's responses align with the job description. If the candidate deviates or provides irrelevant/unwanted answers, gently steer the conversation back by:
  - Asking clarifying or redirecting questions that reflect the job requirements.
  - Reframing the question to emphasize JD-relevant skills, tools, or outcomes.
  - Avoiding tangents that do not contribute to assessing job fit.
 
Conversation so far:
{chat_prompt}
 
Last Candidate Response:
{last_response}
 
Continue the interview with a new technical question or a follow-up that aligns with the job description and addresses any deviation or irrelevance in the candidate's previous response.
"""
        # Get LLM response
        response = llm._call(full_prompt, user)
        reply = response['data']['content']
 
        # Add interviewer response to history
        st.session_state.chat_history.append({"role": "interviewer", "content": reply})
 
        # Advance depth level
        st.session_state.depth_level += 1
 
        # Advance question index
        st.session_state.question_index += 1
 
        # Clear input safely
        st.session_state.temp_input = ""
        st.session_state.response_sent = False
        st.rerun()
 
    # Evaluation option after all questions
    if st.session_state.question_index >= len(st.session_state.questions):
        full_chat = "\n".join(
            f"{msg['role'].capitalize()}: {msg['content']}"
            for msg in st.session_state.chat_history
        )
 
        if st.button("Evaluate Interview"):
            with st.spinner("Evaluating interview..."):
                evaluation = evaluate_interview_transcript(full_chat)
                st.subheader("📊 Interview Evaluation")
                st.text_area("Evaluation Summary", evaluation, height=400)
 
    # Reset chat option
    if st.button("Reset Chat"):
        st.session_state.chat_history = []
        st.session_state.chat_mode_active = False
        st.session_state.temp_input = ""
        st.session_state.response_sent = False
        st.session_state.depth_level = 1
        st.session_state.question_index = 0
        st.session_state.asked_questions = set()
        st.rerun()

        