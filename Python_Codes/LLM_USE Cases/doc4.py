import docx
import pdfplumber
from striprtf.striprtf import rtf_to_text
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

def is_bullet(paragraph: Paragraph):
    return paragraph.style and paragraph.style.name in ['List Bullet', 'List Paragraph']

def extract_docx_content(doc):
    content = {
        'paragraphs': [],
        'tables': [],
        'images': [],
        'headers': [],
        'footers': []
    }

    for section in doc.sections:
        if section.header and section.header.paragraphs:
            content['headers'].append({
                'text': section.header.paragraphs[0].text,
                'alignment': section.header.paragraphs[0].alignment
            })
        if section.footer and section.footer.paragraphs:
            content['footers'].append({
                'text': section.footer.paragraphs[0].text,
                'alignment': section.footer.paragraphs[0].alignment
            })

    for paragraph in doc.paragraphs:
        para_data = {
            'text': paragraph.text,
            'alignment': paragraph.alignment,
            'style': paragraph.style.name,
            'is_bullet': is_bullet(paragraph),
            'runs': []
        }

        for run in paragraph.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_name': run.font.name,
                'font_color': (
                    tuple(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                )
            }
            para_data['runs'].append(run_data)

        content['paragraphs'].append(para_data)

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = {
                    'text': cell.text,
                    'width': getattr(cell, 'width', None),
                }
                row_data.append(cell_data)
            table_data.append(row_data)
        content['tables'].append(table_data)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = {
                'filename': rel.target_ref.split('/')[-1],
                'content': rel._target.blob
            }
            content['images'].append(image_data)

    return content

def read_docx(file_path):
    doc = docx.Document(file_path)
    return extract_docx_content(doc)

def read_pdf(file_path):
    content = {
        'paragraphs': [],
        'tables': [],
        'images': [],  # pdfplumber does not support image extraction
        'headers': [],
        'footers': []
    }

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            width, height = page.width, page.height
            words = page.extract_words()
            lines = page.extract_text().split('\n') if page.extract_text() else []

            for line in lines:
                y = words[0]['top'] if words else 0
                text_data = {
                    'text': line,
                    'alignment': 0,
                    'runs': [{'text': line}]
                }

                if y < 100:
                    content['headers'].append(text_data)
                elif y > height - 100:
                    content['footers'].append(text_data)
                else:
                    content['paragraphs'].append(text_data)

            tables = page.extract_tables()
            for table in tables:
                table_data = []
                for row in table:
                    row_data = [{'text': cell if cell else ''} for cell in row]
                    table_data.append(row_data)
                content['tables'].append(table_data)

    return content

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)

def save_as_docx(content, output_path):
    doc = docx.Document()

    for header_data in content['headers']:
        section = doc.sections[0]
        header = section.header
        if header.paragraphs:
            header_para = header.paragraphs[0]
        else:
            header_para = header.add_paragraph()
        header_para.text = header_data['text']
        if 'alignment' in header_data:
            header_para.alignment = header_data['alignment']

    for para_data in content['paragraphs']:
        if para_data.get('is_bullet'):
            paragraph = doc.add_paragraph(style='List Bullet')
        else:
            paragraph = doc.add_paragraph()

        if 'alignment' in para_data:
            paragraph.alignment = para_data['alignment']

        for run_data in para_data['runs']:
            run = paragraph.add_run(run_data['text'])
            if run_data.get('bold'):
                run.bold = True
            if run_data.get('italic'):
                run.italic = True
            if run_data.get('underline'):
                run.underline = True
            if run_data.get('font_size'):
                run.font.size = Pt(run_data['font_size'])
            if run_data.get('font_name'):
                run.font.name = run_data['font_name']
            if run_data.get('font_color'):
                try:
                    run.font.color.rgb = RGBColor(*run_data['font_color'])
                except Exception:
                    pass

    for table_data in content['tables']:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        set_table_borders(table)
        for i, row in enumerate(table_data):
            for j, cell_data in enumerate(row):
                table.cell(i, j).text = cell_data['text']

    for img_data in content['images']:
        image_stream = io.BytesIO(img_data['content'])
        doc.add_picture(image_stream, width=Inches(6))

    for footer_data in content['footers']:
        section = doc.sections[0]
        footer = section.footer
        if footer.paragraphs:
            footer_para = footer.paragraphs[0]
        else:
            footer_para = footer.add_paragraph()
        footer_para.text = footer_data['text']
        if 'alignment' in footer_data:
            footer_para.alignment = footer_data['alignment']

    doc.save(output_path)

def process_file(file_path, file_type):
    if file_type == 'docx':
        content = read_docx(file_path)
    elif file_type == 'pdf':
        content = read_pdf(file_path)
    else:
        raise ValueError("Unsupported file type")

    base_name = os.path.splitext(file_path)[0]
    save_as_docx(content, base_name + '.converted.docx')
    #save_as_pdf(content, base_name + '.converted.pdf')

# Example usage
file_path = r"D:\OneDrive - Wipro\Desktop\Trainng-Perl_Python\Python_Codes\LLM_USE Cases\formatted_resume-Malaya Ranjan Biswal3.docx"
file_type = 'docx'
process_file(file_path, file_type)

