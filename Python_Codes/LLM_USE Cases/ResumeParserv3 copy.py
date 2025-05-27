from docx import Document
 
def generate_python_code_from_docx(input_file):
    doc = Document(input_file)
    code_lines = ["from docx import Document", "", "doc = Document()", ""]
 
    for para in doc.paragraphs:
        text = para.text.replace("'", "\\'")
        code_lines.append(f"doc.add_paragraph('{text}')")
 
    for table in doc.tables:
        rows = len(table.rows)
        cols = len(table.columns)
        code_lines.append(f"table = doc.add_table(rows={rows}, cols={cols})")
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                text = cell.text.replace("'", "\\'")
                code_lines.append(f"table.cell({i}, {j}).text = '{text}'")
 
    code_lines.append("")
    code_lines.append("doc.save('output.docx')")
 
    return "\n".join(code_lines)
 
# Example usage
input_file = r"D:\OneDrive - Wipro\Desktop\Trainng-Perl_Python\Python_Codes\LLM_USE Cases\formatted_resume-Malaya Ranjan Biswal.docx"
output_file = r"D:\OneDrive - Wipro\Desktop\Trainng-Perl_Python\Python_Codes\LLM_USE Cases\formatted_resume-Malaya Ranjan Biswal3.docx"
# Example usage
print(generate_python_code_from_docx(input_file))
 
 
# Example usage
input_file = r"D:\OneDrive - Wipro\Desktop\Trainng-Perl_Python\Python_Codes\LLM_USE Cases\formatted_resume-Malaya Ranjan Biswal.docx"
output_file = r"D:\OneDrive - Wipro\Desktop\Trainng-Perl_Python\Python_Codes\LLM_USE Cases\formatted_resume-Malaya Ranjan Biswal3.docx"
