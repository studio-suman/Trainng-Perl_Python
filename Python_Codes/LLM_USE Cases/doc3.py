from docx import Document


def generate_formatted_resume(parsed_result: Dict[str, any], output_file: str): # type: ignore
    try:
        
        # Create a new Document
        doc = Document()

        # Add name
        doc.add_heading(parsed_result['name'], level=1)

        # Add contact information
        doc.add_paragraph(f"Email: {parsed_result['email']}")
        doc.add_paragraph(f"Phone: {parsed_result['phone']}")

        # Add summary
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(parsed_result['summary'])

        # Add skills
        doc.add_heading("Skills", level=2)
        skills = "\n".join(parsed_result['skills'])
        doc.add_paragraph(skills)

        # Add experience
        doc.add_heading("Experience", level=2)
        for exp in parsed_result['experience']:
            doc.add_paragraph(exp)

        # Add education
        doc.add_heading("Education", level=2)
        for edu in parsed_result['education']:
            doc.add_paragraph(edu)
        
        # Save the document
        doc.save(output_file)
    except Exception as e:
        print(f"An error occurred while generating the resume: {e}")