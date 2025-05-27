class DocxParser:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = None

    def load_document(self):
        # Load the DOCX document
        from docx import Document
        self.document = Document(self.file_path)

    def extract_paragraphs(self):
        # Extract paragraphs from the document
        paragraphs = []
        for para in self.document.paragraphs:
            paragraphs.append(para.text)
        return paragraphs

    def extract_tables(self):
        # Extract tables from the document
        tables = []
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def extract_styles(self):
        # Extract styles from the document
        styles = {}
        for para in self.document.paragraphs:
            styles[para.text] = para.style.name
        return styles

    def parse(self):
        self.load_document()
        return {
            'paragraphs': self.extract_paragraphs(),
            'tables': self.extract_tables(),
            'styles': self.extract_styles()
        }